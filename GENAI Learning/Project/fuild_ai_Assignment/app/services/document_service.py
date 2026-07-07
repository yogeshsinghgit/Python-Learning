from pathlib import Path
from uuid import uuid4

from docx import Document
from docx.shared import Pt
from loguru import logger

from app.schemas.state import AgentState


OUTPUT_DIR = Path("generated_documents")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


async def generate_document(state: AgentState) -> str:
    logger.info("Generating Word document.")

    document = Document()

    title = document.add_heading(
        state.execution_plan.goal,
        level=1,
    )

    title.runs[0].font.size = Pt(22)

    for section in state.generated_sections:
        heading = document.add_heading(
            section.title,
            level=2,
        )

        heading.runs[0].font.size = Pt(16)

        paragraph = document.add_paragraph()
        paragraph.add_run(section.content)

    filename = f"{uuid4().hex}.docx"

    output_path = OUTPUT_DIR / filename

    document.save(output_path)

    logger.success(f"Document generated: {output_path}")

    return str(output_path)